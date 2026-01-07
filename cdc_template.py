
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_cdc_template(report_period="COP24 Q4"):
    # Create output folder
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)

    # Add timestamp to filename
    date_str = datetime.now().strftime("%Y-%m-%d")
    filename = os.path.join(output_dir, f"cdc_template_{report_period.replace(' ', '_')}_{date_str}.docx")

    doc = Document()

    # Title
    title = doc.add_paragraph(f"Zim-TTECH {report_period} DATIM File Narrative")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(16)

    # Section: Number of facilities reporting
    doc.add_heading("Number of facilities reporting", level=1)
    doc.add_paragraph(
        f"In {report_period}, a total of {{total_active_facilities}} facilities were active on the full IMPILO E-HR system. "
        "Raw data files were available from {{raw_data_sites}} sites, of which {{collected_manually}} were collected manually, "
        "{{pushed_via_pipeline}} were pushed via the data pipeline and {{mobile_backups}} were complete backups obtained via the mobile backup."
    )
    doc.add_paragraph("For the period under review, data from {{ehr_facilities_analyzed}} facilities ({{ehr_facilities_analyzed_pct}}%) were successfully analysed.")
    doc.add_paragraph("{{figure1}}")
    doc.add_paragraph("Figure 1: Summary of IMPILO E-HR data flow cascade.", style="Caption")

    # Section: Indicators reported
    doc.add_heading("Number of indicators reported", level=1)
    doc.add_paragraph("{{table1_indicators}}")
    doc.add_paragraph("Table 1: Indicators reported through IMPILO E-HR", style="Caption")

    # Section: Data Quality Metrics
    doc.add_heading("Data Quality Metrics", level=1)
    doc.add_paragraph(
        "In this period TX_CURR, TX_ML, TX_TB, HTS_TST/POS and PMTCT_STAT had the highest proportion of facilities reporting "
        "({{prop_tx_curr}}%, {{prop_tx_ml}}%, {{prop_tx_tb}}%, {{prop_hts_tst}}%, {{prop_pmtct_stat}}%) as shown in Figure 2."
    )
    doc.add_paragraph("{{figure2}}")
    doc.add_paragraph("Figure 2: Proportion of sites reporting individual indicators.", style="Caption")

    doc.add_paragraph("Reporting Consistency")
    doc.add_paragraph("{{figure3}}")
    doc.add_paragraph("Figure 3: Trends in sites successfully reporting.", style="Caption")

    # Section: Challenges
    doc.add_heading("Challenges", level=1)
    doc.add_paragraph("• Database collection")
    doc.add_paragraph("{{challenge_database_collection}}")
    doc.add_paragraph("• Data extraction")
    doc.add_paragraph("{{challenge_data_extraction}}")

    # Section: Remedial Actions
    doc.add_heading("Remedial Actions Taken and Recommendations", level=1)
    doc.add_paragraph("• {{remedial_point_1}}")
    doc.add_paragraph("• {{remedial_point_2}}")
    doc.add_paragraph("• {{remedial_point_3}}")

    # Section: Optimized Sites Data Concordance Analysis
    doc.add_heading("Optimized Sites Data Concordance Analysis", level=1)
    doc.add_paragraph("{{table2_concordance_facility}}")
    doc.add_paragraph("Table 2: MRF/E-HR Data Concordance Analysis for Optimized sites", style="Caption")

    # Section: Overall Data Concordance Analysis
    doc.add_heading("Overall Data Concordance Analysis", level=1)
    table3 = doc.add_table(rows=5, cols=4)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = "Indicator"
    hdr_cells[1].text = "MRF/DHIS2"
    hdr_cells[2].text = "E-HR"
    hdr_cells[3].text = "Concordance"
    for cell in hdr_cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    rows_data = [
        ["HTS_TST", "{{overall_hts_tst_mrf}}", "{{overall_hts_tst_ehr}}", "{{overall_hts_tst_conc}}%"],
        ["HTS_POS", "{{overall_hts_pos_mrf}}", "{{overall_hts_pos_ehr}}", "{{overall_hts_pos_conc}}%"],
        ["TX_NEW", "{{overall_tx_new_mrf}}", "{{overall_tx_new_ehr}}", "{{overall_tx_new_conc}}%"],
        ["TX_CURR", "{{overall_tx_curr_mrf}}", "{{overall_tx_curr_ehr}}", "{{overall_tx_curr_conc}}%"]
    ]
    for r_idx, rdata in enumerate(rows_data, start=1):
        row_cells = table3.rows[r_idx].cells
        for c_idx, val in enumerate(rdata):
            row_cells[c_idx].text = val

    # Section: TX_CURR Concordance Analysis for Harare and Bulawayo
    doc.add_heading("TX_CURR Concordance Analysis for Harare and Bulawayo", level=1)
    doc.add_paragraph("{{figure4}}")
    doc.add_paragraph("Figure 4: Concordance analysis for Harare and Bulawayo sites", style="Caption")

    # Save template
    doc.save(filename)
    print(f"Template created: {filename}")

# Example usage:
create_cdc_template(report_period="COP25 Q1")
