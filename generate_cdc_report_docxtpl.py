#!/usr/bin/env python3
"""
Generate CDC Report Figures 1, 2, 4 and Tables 1, 2, 3.
Complete Final Version.
"""

import argparse
import os
from datetime import datetime
import pandas as pd
import matplotlib
matplotlib.use('Agg')  
import matplotlib.pyplot as plt
import seaborn as sns
import matplotlib.patches as mpatches
from matplotlib.lines import Line2D

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------
# 1. CONSTANTS & MAPPINGS
# ---------------

TARGETS = {
    'TX_CURR': 823, 'TX_TB': 823, 'TX_ML': 823, 'HTS_TST and HTS_POS': 823,
    'PMTCT_STAT': 761, 'PMTCT_ART': 659, 'TB_PREV': 657, 'TX_PVLS': 823,
    'HTS_INDEX': 823, 'TX_NEW': 823, 'TB_ART': 657, 'PREP_NEW': 628,
    'PMTCT_FO': 661, 'HTS_SELF': 823, 'TB_STAT': 657, 'PMTCT_EID': 655,
    'CXCA_SCRN': 628, 'CXCA_TX': 473, 'PREP_CT': 628, 'TX_RTT': 823,
    'PMTCT_HEI': 661,
}

MAP_RULES = {
    'TX_CURR': ['TX_CURR TA'], 'TX_TB': ['TX_TB(DENOM) TA', 'TX_TB TA'],
    'TX_ML': ['TX_ML TA'], 'HTS_TST': ['HTS_TST TA'], 'HTS_POS': ['HTS_POS TA'],
    'PMTCT_STAT': ['PMTCT_STAT(DENOM) TA', 'PMTCT_STAT TA'], 'PMTCT_ART': ['PMTCT_ART TA'],
    'TB_PREV': ['TB_PREV(DENOM) TA', 'TB_PREV TA'], 'TX_PVLS': ['TX_PVLS(DENOM) TA', 'TX_PVLS TA'],
    'HTS_INDEX': ['HTS_INDEX TA'], 'TX_NEW': ['TX_NEW TA'], 'TB_ART': ['TB_ART TA'],
    'PREP_NEW': ['PrEP_NEW TA', 'PREP_NEW TA'], 'PMTCT_FO': ['PMTCT_FO TA'],
    'HTS_SELF': ['HTS_SELF TA'], 'TB_STAT': ['TB_STAT(DENOM) TA', 'TB_STAT TA'],
    'PMTCT_EID': ['PMTCT_EID TA'], 'CXCA_SCRN': ['CXCA_SCRN TA'], 'CXCA_TX': ['CXCA_TX TA'],
    'PREP_CT': ['PrEP_CT TA', 'PREP_CT TA'], 'TX_RTT': ['TX_RTT TA'], 'PMTCT_HEI': ['PMTCT_HEI TA'],
}

INDICATOR_ORDER = [
    'TX_CURR', 'TX_TB', 'TX_ML', 'HTS_TST and HTS_POS', 'PMTCT_STAT', 'PMTCT_ART',
    'TB_PREV', 'TX_PVLS', 'HTS_INDEX', 'TX_NEW', 'TB_ART', 'PREP_NEW', 'PMTCT_FO',
    'HTS_SELF', 'TB_STAT', 'PMTCT_EID', 'CXCA_SCRN', 'CXCA_TX', 'PREP_CT', 'TX_RTT', 'PMTCT_HEI'
]

CONCORDANCE_MAP = {
    "HTS_TST":  ("MRF HTS_TX_NEW (Total Tests)",        "DATIM HTS_TX_NEW (Total Tests)"),
    "HTS_POS":  ("MRF HTS_TX_NEW (Total Positives)",    "DATIM HTS_TX_NEW (Total Positives)"),
    "TX_NEW":   ("MRF HTS_TX_NEW (Total Initiations)",  "DATIM HTS_TX_NEW (Total Initiations)"),
    "TX_CURR":  ("MRF TX_CURR",                         "DATIM TX_CURR"),
}

# ---------------
# 2. HELPER FUNCTIONS
# ---------------

def load_consistency(path: str) -> pd.DataFrame:
    df = pd.read_excel(path, engine='openpyxl', sheet_name=0)
    df.columns = [str(c).strip() for c in df.columns]
    df = df.loc[:, ~df.columns.duplicated()]
    ind_col = None; num_col = None
    for c in df.columns:
        cname = str(c).strip().lower()
        if 'indicator' in cname: ind_col = c
        if ('number' in cname or 'no.' in cname) and 'report' in cname: num_col = c
    if ind_col is None or num_col is None: ind_col = df.columns[0]; num_col = df.columns[1]
    return pd.DataFrame({
        'Indicator': df[ind_col].astype(str).str.strip(),
        'ReportingUnits': pd.to_numeric(df[num_col], errors='coerce').fillna(0).astype(int)
    })

def pick_reporting(cons_df: pd.DataFrame, patterns: list[str]) -> int:
    for pat in patterns:
        row = cons_df[cons_df['Indicator'] == pat]
        if not row.empty: return int(row.iloc[0]['ReportingUnits'])
    return 0

def build_21_indicator_table(cons_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    values = {}
    for key, pats in MAP_RULES.items(): values[key] = pick_reporting(cons_df, pats)
    hts_combined = max(values.get('HTS_TST', 0), values.get('HTS_POS', 0))
    for ind in INDICATOR_ORDER:
        reporting = hts_combined if ind == 'HTS_TST and HTS_POS' else values.get(ind, 0)
        tgt = TARGETS[ind]
        comp = round(100 * reporting / tgt, 1) if tgt else 0.0
        rows.append({'Indicator': ind, 'Number_of_facilities_reporting': reporting, 'Target': tgt, 'Reporting_Completeness_%': comp})
    return pd.DataFrame(rows)

def build_figure1_cascade(total, manual, pipeline, mobile, analyzed, out_png):
    """
    Generate Detailed Organogram for Figure 1.
    """
    # --- Calculations ---
    total_collected = manual + pipeline + mobile
    not_collected = total - total_collected
    
    coll_pct = (total_collected / total * 100) if total > 0 else 0
    
    # Estimate success/fail (since we only have global analyzed count)
    total_failed = total_collected - analyzed
    analyzed_pct = (analyzed / total_collected * 100) if total_collected > 0 else 0
    failed_pct = (total_failed / total_collected * 100) if total_collected > 0 else 0
    
    # Split outcomes proportionally
    def split_outcomes(source_total):
        if total_collected == 0: return 0, 0
        succ = int(source_total * (analyzed / total_collected))
        fail = source_total - succ
        return succ, fail

    man_succ, man_fail = split_outcomes(manual)
    mob_succ, mob_fail = split_outcomes(mobile)
    pipe_succ, pipe_fail = split_outcomes(pipeline)
    
    # Fix rounding
    curr_succ = man_succ + mob_succ + pipe_succ
    man_succ += (analyzed - curr_succ)
    
    def fmt(val, parent_val):
        pct = (val / parent_val * 100) if parent_val > 0 else 0
        return f"{int(val)} ({pct:.0f}%)"

    # --- Plotting ---
    fig, ax = plt.subplots(figsize=(16, 14))
    ax.set_xlim(0, 100); ax.set_ylim(0, 100); ax.axis('off')
    
    y_L1=96; y_L2=84; y_L3=62; y_L4=32; y_L5=8
    xy_total=(50, y_L1); xy_coll=(30, y_L2); xy_not=(70, y_L2)
    xy_src_man=(16, y_L3); xy_src_mob=(50, y_L3); xy_src_dpl=(84, y_L3)
    xy_man_S=(8, y_L4); xy_man_F=(24, y_L4); xy_mob_S=(42, y_L4); xy_mob_F=(58, y_L4); xy_dpl_S=(76, y_L4); xy_dpl_F=(92, y_L4)
    
    c_blue='#A0CBE8'; c_green='#EBF7EB'; c_red='#FADBD8'; c_src='#FEF5E7'; c_line='#555555'

    def connect(xy1, xy2): ax.plot([xy1[0], xy2[0]], [xy1[1], xy2[1]], color=c_line, lw=1.5, zorder=0)
    def add_box(xy, text, bg_color, fontsize=10, boxstyle="round,pad=1"):
        ax.text(xy[0], xy[1], text, ha='center', va='center', fontsize=fontsize, fontweight='bold', bbox=dict(boxstyle=boxstyle, fc=bg_color, ec="black", lw=1))

    # Lines
    connect(xy_total, xy_coll); connect(xy_total, xy_not)
    connect(xy_coll, xy_src_man); connect(xy_coll, xy_src_mob); connect(xy_coll, xy_src_dpl)
    connect(xy_src_man, xy_man_S); connect(xy_src_man, xy_man_F)
    connect(xy_src_mob, xy_mob_S); connect(xy_src_mob, xy_mob_F)
    connect(xy_src_dpl, xy_dpl_S); connect(xy_src_dpl, xy_dpl_F)
    
    # Boxes
    add_box(xy_total, f"Total Full Impilo EHR Sites\n{total}", c_blue, fontsize=12)
    add_box(xy_coll, f"Backups Collected\n{fmt(total_collected, total)}", c_blue)
    add_box(xy_not, f"Backups NOT Collected\n{fmt(not_collected, total)}", '#F5B7B1')
    
    add_box(xy_src_man, f"Physical Backups\nCollected\n{fmt(manual, total_collected)}", c_src)
    add_box(xy_src_mob, f"Mobile Application\nData Collection\n{fmt(mobile, total_collected)}", c_src)
    add_box(xy_src_dpl, f"DPL Data\nCollection\n{fmt(pipeline, total_collected)}", c_src)
    
    s_lbl = "Extracted and available for\nanalysis and reporting"
    f_lbl = "Backups that failed\nto process"
    
    add_box(xy_man_S, f"{s_lbl}\n{fmt(man_succ, manual)}", c_green, fontsize=8, boxstyle="square,pad=0.8")
    add_box(xy_man_F, f"{f_lbl}\n{fmt(man_fail, manual)}", c_red, fontsize=8, boxstyle="square,pad=0.8")
    add_box(xy_mob_S, f"{s_lbl}\n{fmt(mob_succ, mobile)}", c_green, fontsize=8, boxstyle="square,pad=0.8")
    add_box(xy_mob_F, f"{f_lbl}\n{fmt(mob_fail, mobile)}", c_red, fontsize=8, boxstyle="square,pad=0.8")
    add_box(xy_dpl_S, f"{s_lbl}\n{fmt(pipe_succ, pipeline)}", c_green, fontsize=8, boxstyle="square,pad=0.8")
    add_box(xy_dpl_F, f"{f_lbl}\n{fmt(pipe_fail, pipeline)}", c_red, fontsize=8, boxstyle="square,pad=0.8")
    
    summary_text = (f"SUMMARY STATISTICS\n────────────────────────────────────────\n"
                    f"Collection Rate: {coll_pct:.1f}% ({total_collected}/{total})\n"
                    f"Total Successfully Extracted: {analyzed_pct:.1f}% ({analyzed}/{total_collected})\n"
                    f"Total Failed to Process: {failed_pct:.1f}% ({total_failed}/{total_collected})")
    ax.text(50, y_L5, summary_text, ha='center', va='center', fontsize=12, fontfamily='monospace', bbox=dict(boxstyle="round,pad=1.2", fc='#FAFAFA', ec='black', lw=1.5))

    plt.title('Figure 1: Summary of IMPILO E-HR Data Flow Cascade', fontsize=16, fontweight='bold', y=0.99)
    plt.tight_layout(); plt.savefig(out_png, dpi=200); plt.close()

def build_figure2(table_df: pd.DataFrame, out_png: str, period_label: str):
    sns.set(style='whitegrid')
    plt.figure(figsize=(12,8))
    plot_df = table_df.sort_values('Reporting_Completeness_%', ascending=False)
    ax = sns.barplot(x='Indicator', y='Reporting_Completeness_%', data=plot_df, color='#4C78A8')
    plt.ylabel('Reporting Completeness (%)'); plt.xlabel('Indicator'); plt.ylim(0,100)
    plt.title(f'Figure 2: Proportion of sites reporting individual indicators in IMPILO E-HR {period_label}')
    plt.xticks(rotation=45, ha='right')
    for i, p in enumerate(ax.patches):
        pct = plot_df.iloc[i]['Reporting_Completeness_%']
        ax.annotate(f"{pct}%", (p.get_x()+p.get_width()/2, p.get_height()), ha='center', va='bottom', fontsize=9)
    plt.tight_layout(); plt.savefig(out_png, dpi=200); plt.close()

def concordance_value(mrf, ehr):
    if mrf == 0 and ehr == 0: return 100.0
    if mrf == 0 and ehr > 0: return 0.0
    return max(0.0, min(100.0, (1 - abs(ehr - mrf) / mrf) * 100))

def load_concordance(path: str) -> pd.DataFrame:
    raw = pd.read_excel(path, engine='openpyxl')
    raw.columns = [str(c).strip() for c in raw.columns]
    rows = []
    for ind, (mrf_col, ehr_col) in CONCORDANCE_MAP.items():
        if mrf_col not in raw.columns or ehr_col not in raw.columns: continue
        tmp = raw[["Facility", "Province", "District", mrf_col, ehr_col]].copy()
        tmp["Indicator"] = ind
        tmp["MRF"] = pd.to_numeric(tmp[mrf_col], errors="coerce").fillna(0)
        tmp["EHR"] = pd.to_numeric(tmp[ehr_col], errors="coerce").fillna(0)
        tmp["Concordance_%"] = tmp.apply(lambda r: concordance_value(r["MRF"], r["EHR"]), axis=1)
        rows.append(tmp[["Facility", "Province", "District", "Indicator", "MRF", "EHR", "Concordance_%"]])
    return pd.concat(rows, ignore_index=True) if rows else pd.DataFrame()

def compute_table3_context(df: pd.DataFrame) -> dict:
    context = {}
    required_cols = {'Indicator', 'MRF', 'EHR'}; indicators = ['HTS_TST', 'HTS_POS', 'TX_NEW', 'TX_CURR']
    if not required_cols.issubset(df.columns): return context
    for ind in indicators:
        sub = df[df['Indicator'].astype(str).str.upper() == ind]
        mrf_sum = sub['MRF'].sum() if not sub.empty else 0
        ehr_sum = sub['EHR'].sum() if not sub.empty else 0
        conc = round((ehr_sum / mrf_sum * 100), 1) if mrf_sum > 0 else 0.0
        context[f"overall_{ind.lower()}_mrf"] = f"{int(mrf_sum):,}"
        context[f"overall_{ind.lower()}_ehr"] = f"{int(ehr_sum):,}"
        context[f"overall_{ind.lower()}_conc"] = str(conc)
    return context

def build_figure4(df: pd.DataFrame, out_png: str):
    required_cols = {'Indicator', 'Province', 'MRF', 'EHR'}
    if not required_cols.issubset(df.columns): return
    target_provinces = ['Harare', 'Bulawayo']
    mask_prov = df['Province'].astype(str).str.strip().apply(lambda x: any(p.lower() in x.lower() for p in target_provinces))
    mask_ind = df['Indicator'].astype(str).str.upper() == 'TX_CURR'
    subset = df[mask_prov & mask_ind].copy()
    if subset.empty: return

    def normalize_name(p):
        p_str = str(p).lower()
        if 'harare' in p_str: return 'Harare'
        if 'bulawayo' in p_str: return 'Bulawayo'
        return p
    subset['Province'] = subset['Province'].apply(normalize_name)
    grouped = subset.groupby('Province')[['MRF', 'EHR']].sum().reset_index()
    total_row = pd.DataFrame({'Province': ['Total'], 'MRF': [grouped['MRF'].sum()], 'EHR': [grouped['EHR'].sum()]})
    ordered_provinces = pd.concat([total_row, grouped], ignore_index=True)
    plot_df = ordered_provinces.melt(id_vars='Province', value_vars=['MRF', 'EHR'], var_name='Source', value_name='Count')
    ordered_provinces['Concordance'] = ordered_provinces.apply(lambda r: (r['EHR'] / r['MRF'] * 100) if r['MRF'] > 0 else 0, axis=1)
    x_order = ['Total'] + [p for p in target_provinces if p in grouped['Province'].unique()]

    sns.set(style='white'); fig, ax1 = plt.subplots(figsize=(11, 6))
    custom_colors = {'MRF': '#104E8B', 'EHR': '#00FF00'} 
    sns.barplot(x='Province', y='Count', hue='Source', data=plot_df, order=x_order, palette=custom_colors, edgecolor="black", ax=ax1)
    if ax1.get_legend(): ax1.get_legend().remove()
    ax1.set_ylabel('Number of clients on ART', fontsize=12); ax1.set_xlabel('Province', fontsize=12)
    ax1.set_title('TX_CURR Concordance: MRF vs EHR', fontsize=14, pad=20)
    for container in ax1.containers: ax1.bar_label(container, fmt='{:,.0f}', padding=3, fontsize=10)

    ax2 = ax1.twinx(); ax2.set_ylim(0, 115)
    y_values = [ordered_provinces.loc[ordered_provinces['Province'] == region, 'Concordance'].values[0] for region in x_order]
    x_coords = range(len(x_order))
    ax2.plot(x_coords, y_values, color='#D62728', marker='D', markersize=8, linestyle='None')
    for x, y in zip(x_coords, y_values): ax2.text(x, y + 4, f"{y:.1f}%", color='#D62728', ha='center', fontweight='bold', fontsize=11)
    ax2.set_yticks([]); ax2.set_ylabel(''); ax2.spines['right'].set_visible(False); ax2.spines['top'].set_visible(False)
    ax1.spines['top'].set_visible(False); ax1.spines['right'].set_visible(False)
    legend_elements = [mpatches.Patch(facecolor='#104E8B', edgecolor='black', label='MRF'), mpatches.Patch(facecolor='#00FF00', edgecolor='black', label='EHR'), Line2D([0], [0], color='#D62728', marker='D', linestyle='None', markersize=8, label='Concordance')]
    ax1.legend(handles=legend_elements, title='Key', bbox_to_anchor=(1.02, 1), loc='upper left')
    ax1.grid(True, axis='y', alpha=0.3); plt.tight_layout(); plt.savefig(out_png, dpi=200, bbox_inches='tight'); plt.close()

def aggregate_table3_data(df: pd.DataFrame) -> pd.DataFrame:
    indicators = ['HTS_TST', 'HTS_POS', 'TX_NEW', 'TX_CURR']; data = []
    for ind in indicators:
        sub_df = df[df['Indicator'].astype(str).str.upper() == ind]
        mrf_sum = sub_df['MRF'].sum() if not sub_df.empty else 0
        ehr_sum = sub_df['EHR'].sum() if not sub_df.empty else 0
        conc = (ehr_sum / mrf_sum * 100) if mrf_sum > 0 else (100.0 if ehr_sum == 0 else 0.0)
        data.append({"Indicator": ind, "MRF/DHIS2": int(mrf_sum), "E-HR": int(ehr_sum), "Concordance": conc})
    return pd.DataFrame(data)

def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), color_hex); tcPr.append(shd)
def get_concordance_color(val): return 'C6EFCE' if val >= 95 else ('FFEB9C' if val >= 90 else 'F8CBAD')

def build_overall_table_subdoc(tpl: DocxTemplate, agg_df: pd.DataFrame):
    sub = tpl.new_subdoc(); tbl = sub.add_table(rows=1, cols=4)
    try: tbl.style = 'Table Grid'
    except KeyError: pass
    headers = ["Indicator", "MRF/DHIS2", "E-HR", "Concordance"]
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers): hdr_cells[i].text = h; [run.bold for p in hdr_cells[i].paragraphs for run in p.runs]
    for _, row in agg_df.iterrows():
        row_cells = tbl.add_row().cells
        row_cells[0].text = str(row["Indicator"]); row_cells[1].text = f"{int(row['MRF/DHIS2']):,}"
        row_cells[2].text = f"{int(row['E-HR']):,}"; conc = row["Concordance"]; row_cells[3].text = f"{conc:.1f}%"
        if (c := get_concordance_color(conc)): shade_cell(row_cells[3], c)
    return sub

def build_table2_subdoc(tpl: DocxTemplate, df: pd.DataFrame):
    sub = tpl.new_subdoc(); cols = 7; tbl = sub.add_table(rows=1, cols=cols)
    try: tbl.style = 'Table Grid'
    except KeyError: pass
    headers = ["Facility", "Province", "District", "Indicator", "MRF", "EHR", "Concordance_%"]
    hdr_cells = tbl.rows[0].cells
    for i, col in enumerate(headers): hdr_cells[i].text = str(col); [run.bold for p in hdr_cells[i].paragraphs for run in p.runs]
    for _, row_series in df.iterrows():
        row_cells = tbl.add_row().cells; row_data = [row_series.get(h, "") for h in headers]
        for i, val in enumerate(row_data):
            txt = "" if pd.isna(val) else str(val)
            if isinstance(val, float) and val.is_integer(): txt = str(int(val))
            elif isinstance(val, float) and headers[i] == "Concordance_%": txt = f"{val:.1f}%"
            row_cells[i].text = txt
            if headers[i] == "Concordance_%":
                try:
                    if (c := get_concordance_color(float(val))): shade_cell(row_cells[i], c)
                    row_cells[i].text = f"{float(val):.1f}%"
                except (ValueError, TypeError): pass
    return sub

def build_table1_subdoc(tpl: DocxTemplate, table_df: pd.DataFrame):
    sub = tpl.new_subdoc(); tbl = sub.add_table(rows=1, cols=4)
    headers = ['Indicator', 'Number of facilities reporting', 'Target', 'Reporting Completeness']
    for j, h in enumerate(headers): cell = tbl.rows[0].cells[j]; cell.text = h; [run.bold for p in cell.paragraphs for run in p.runs]
    for _, r in table_df.iterrows():
        row = tbl.add_row().cells
        row[0].text = str(r['Indicator']); row[1].text = str(int(r['Number_of_facilities_reporting']))
        row[2].text = str(int(r['Target'])); row[3].text = f"{r['Reporting_Completeness_%']}%"
    return sub

# ---------------
# 3. MAIN EXECUTION
# ---------------
def main():
    parser = argparse.ArgumentParser(description='Generate CDC Report with Figures 1, 2, 4.')
    parser.add_argument('--consistency', required=True, help='Path to Consistency Report')
    parser.add_argument('--concordance', help='Path to Concordance Analysis')
    parser.add_argument('--template', default='cdc_template.docx', help='Template path')
    parser.add_argument('--out', default=None, help='Output path')
    parser.add_argument('--period', default='COP24 Q4', help='Period label')
    parser.add_argument('--outdir', default='output', help='Output directory')
    # Args for Figure 1 & Text
    parser.add_argument('--total_sites', type=int, default=0)
    parser.add_argument('--raw_sites', type=int, default=0)
    parser.add_argument('--analyzed_sites', type=int, default=0)
    parser.add_argument('--manual', type=int, default=0)
    parser.add_argument('--pipeline', type=int, default=0)
    parser.add_argument('--mobile', type=int, default=0)
    args = parser.parse_args()

    os.makedirs(args.outdir, exist_ok=True)
    stamp = datetime.now().strftime('%Y-%m-%d')
    print("Starting generation...")

    # 1. Consistency
    cons_df = load_consistency(args.consistency)
    table_df = build_21_indicator_table(cons_df)
    csv_path = os.path.join(args.outdir, f'figure2_table1_{args.period.replace(" ", "")}_{stamp}.csv')
    table_df.to_csv(csv_path, index=False)

    # 2. Figure 1 (Cascade)
    fig1_path = os.path.join(args.outdir, f'figure1_cascade_{args.period.replace(" ", "")}_{stamp}.png')
    build_figure1_cascade(args.total_sites, args.manual, args.pipeline, args.mobile, args.analyzed_sites, fig1_path)

    # 3. Figure 2 (Completeness)
    fig2_path = os.path.join(args.outdir, f'figure2_reporting_completeness_{args.period.replace(" ", "")}_{stamp}.png')
    build_figure2(table_df, fig2_path, args.period)

    # 4. Doc Generation
    tpl = DocxTemplate(args.template)
    sub_table1 = build_table1_subdoc(tpl, table_df)
    sub_table2 = ""; table3_context = {}; fig4_path = ""
    
    if args.concordance:
        conc_df = load_concordance(args.concordance)
        t2_csv = os.path.join(args.outdir, f'table2_concordance_{args.period.replace(" ", "")}_{stamp}.csv')
        conc_df.to_csv(t2_csv, index=False)
        table3_df = aggregate_table3_data(conc_df)
        t3_csv = os.path.join(args.outdir, f'table3_overall_concordance_{args.period.replace(" ", "")}_{stamp}.csv')
        table3_df.to_csv(t3_csv, index=False)
        sub_table2 = build_table2_subdoc(tpl, conc_df)
        table3_context = compute_table3_context(conc_df)
        fig4_path = os.path.join(args.outdir, f'figure4_concordance_{args.period.replace(" ", "")}_{stamp}.png')
        build_figure4(conc_df, fig4_path)

    analyzed_pct = round(args.analyzed_sites / args.total_sites * 100, 1) if args.total_sites > 0 else 0
    context = {
        'total_active_facilities': args.total_sites,
        'raw_data_sites': args.raw_sites,
        'collected_manually': args.manual, 'pushed_via_pipeline': args.pipeline, 'mobile_backups': args.mobile,
        'ehr_facilities_analyzed': args.analyzed_sites, 'ehr_facilities_analyzed_pct': analyzed_pct,
        'figure1': InlineImage(tpl, fig1_path, width=Mm(140)),
        'figure2': InlineImage(tpl, fig2_path, width=Mm(140)),
        'table1_indicators': sub_table1, 'table2_concordance_facility': sub_table2,
        'figure4': '', **table3_context
    }
    if fig4_path and os.path.exists(fig4_path): context['figure4'] = InlineImage(tpl, fig4_path, width=Mm(140))
        
    tpl.render(context)
    out_docx = args.out or os.path.join(args.outdir, f'CDC_Report_{args.period.replace(" ", "")}_{stamp}.docx')
    tpl.save(out_docx)
    print(f'Done! Report saved to: {out_docx}')

if __name__ == '__main__':
    main()