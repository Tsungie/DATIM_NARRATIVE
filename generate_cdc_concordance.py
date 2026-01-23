
#!/usr/bin/env python3
"""
Generate Table 2 (facility-level concordance), Table 3 (overall concordance by indicator),
and Figure 4 (Harare vs Bulawayo concordance) from Concordance_Analysis_Complete.xlsx,
then inject all three into cdc_template.docx.

Usage:
  python generate_cdc_concordance.py \
    --excel "Concordance_Analysis_Complete.xlsx" \
    --template "cdc_template.docx" \
    --outdir output \
    --period "COP24 Q4" \
    --out output/CDC_Concordance_COP24Q4_<YYYY-MM-DD>.docx

Dependencies:
  pip install pandas numpy matplotlib seaborn docxtpl python-docx openpyxl docxcompose
"""

import argparse
import os
from datetime import datetime
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# -----------------------------
# Indicator mapping (from your file)
# -----------------------------
INDICATOR_PAIR_MAP = {
    "HTS_TST":  ("MRF HTS_TX_NEW (Total Tests)",        "DATIM HTS_TX_NEW (Total Tests)"),
    "HTS_POS":  ("MRF HTS_TX_NEW (Total Positives)",    "DATIM HTS_TX_NEW (Total Positives)"),
    "TX_NEW":   ("MRF HTS_TX_NEW (Total Initiations)",  "DATIM HTS_TX_NEW (Total Initiations)"),
    "TX_CURR":  ("MRF TX_CURR",                         "DATIM TX_CURR"),
}

HB_PROVINCES = {"Harare Metropolitan", "Bulawayo Metropolitan"}

# -----------------------------
# Concordance functions (CDC-consistent)
# -----------------------------
def concordance_point(mrf, ehr):
    """
    Facility-level concordance (%), CDC-style:
      - If MRF=0 and EHR=0 => 100
      - If MRF=0 and EHR>0 => 0
      - Else: (1 - |EHR - MRF| / MRF) * 100, clipped to [0, 100]
    """
    if pd.isna(mrf) or pd.isna(ehr):
        return np.nan
    if mrf == 0 and ehr == 0:
        return 100.0
    if mrf == 0 and ehr > 0:
        return 0.0
    pct = (1.0 - abs(ehr - mrf) / mrf) * 100.0
    return float(max(0.0, min(100.0, pct)))

def concordance_overall(mrf_series, ehr_series):
    """
    Overall (aggregate) concordance (%), CDC-style:
      (1 - sum(|EHR_i - MRF_i|) / sum(MRF_i)) * 100  if sum(MRF) > 0
      else 100 when all zeros.
    """
    mrf_sum = mrf_series.sum()
    if mrf_sum == 0:
        return 100.0
    diff_sum = (ehr_series - mrf_series).abs().sum()
    pct = (1.0 - diff_sum / mrf_sum) * 100.0
    return float(max(0.0, min(100.0, pct)))

# -----------------------------
# Styling helpers for color-coding
# -----------------------------
def shade_cell(cell, color_hex):
    # color_hex: 'C6EFCE' (green), 'FFEB9C' (amber), 'F8CBAD' (red)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def color_for_conc(pct):
    if pd.isna(pct):
        return None
    p = float(pct)
    if p >= 95.0:
        return 'C6EFCE'  # green
    if 90.0 <= p < 95.0:
        return 'FFEB9C'  # amber
    return 'F8CBAD'      # red

# -----------------------------
# Core build functions
# -----------------------------
def load_concordance_excel(path):
    df = pd.read_excel(path, engine="openpyxl")
    # Normalize columns
    df.columns = [str(c).strip() for c in df.columns]
    # Ensure expected ID columns exist
    for col in ["Facility", "Province", "District"]:
        if col not in df.columns:
            raise ValueError(f"Required column missing: {col}")
    return df

def build_table2(df_raw):
    """
    Build facility-level concordance for the 4 indicators:
    returns DataFrame with columns:
    ['Facility','Province','District','Indicator','MRF','E_HR','Concordance_%']
    """
    rows = []
    for ind, (mrf_col, ehr_col) in INDICATOR_PAIR_MAP.items():
        missing = [c for c in (mrf_col, ehr_col) if c not in df_raw.columns]
        if missing:
            raise ValueError(f"Missing expected columns for {ind}: {missing}")

        tmp = df_raw[["Facility", "Province", "District", mrf_col, ehr_col]].copy()
        tmp["Indicator"] = ind
        tmp["MRF"] = pd.to_numeric(tmp[mrf_col], errors="coerce").fillna(0).astype(float)
        tmp["E_HR"] = pd.to_numeric(tmp[ehr_col], errors="coerce").fillna(0).astype(float)
        tmp["Concordance_%"] = tmp.apply(lambda r: concordance_point(r["MRF"], r["E_HR"]), axis=1)

        rows.append(tmp[["Facility", "Province", "District", "Indicator", "MRF", "E_HR", "Concordance_%"]])

    table2 = pd.concat(rows, ignore_index=True)
    return table2

def build_table3(table2):
    """
    Aggregate facility-level to overall by indicator using CDC-style overall formula.
    Returns DataFrame with columns:
    ['Indicator','MRF_Total','E_HR_Total','Overall_Concordance_%']
    """
    out = []
    for ind, grp in table2.groupby("Indicator"):
        mrf_sum = grp["MRF"].sum()
        ehr_sum = grp["E_HR"].sum()
        conc = concordance_overall(grp["MRF"], grp["E_HR"])
        out.append({"Indicator": ind,
                    "MRF_Total": int(round(mrf_sum)),
                    "E_HR_Total": int(round(ehr_sum)),
                    "Overall_Concordance_%": round(conc, 1)})
    return pd.DataFrame(out)

def build_figure4(table2, out_png, period_label):
    """
    Figure 4: Harare vs Bulawayo overall concordance by indicator
    """
    hb = table2[table2["Province"].isin(HB_PROVINCES)].copy()
    if hb.empty:
        print("[WARN] No rows for Harare/Bulawayo in data; creating empty figure.")
        plt.figure(figsize=(8,4))
        plt.text(0.5,0.5,"No Harare/Bulawayo data", ha="center")
        plt.axis('off'); plt.tight_layout(); plt.savefig(out_png, dpi=200); plt.close()
        return

    rows = []
    for (prov, ind), grp in hb.groupby(["Province","Indicator"]):
        conc = concordance_overall(grp["MRF"], grp["E_HR"])
        rows.append({"Province": prov, "Indicator": ind, "Concordance_%": round(conc, 1)})
    figdf = pd.DataFrame(rows)

    sns.set(style="whitegrid")
    plt.figure(figsize=(12, 6))
    ax = sns.barplot(data=figdf, x="Indicator", y="Concordance_%", hue="Province")
    ax.set_ylabel("Concordance (%)")
    ax.set_xlabel("Indicator")
    ax.set_title(f"Concordance analysis for Harare and Bulawayo sites — {period_label}")
    ax.set_ylim(0, 100)
    plt.xticks(rotation=45, ha="right")
    plt.legend(title="Province")
    plt.tight_layout()
    plt.savefig(out_png, dpi=200)
    plt.close()

# -----------------------------
# Subdocument builders (docxtpl)
# -----------------------------
def build_table2_subdoc(tpl, table2_df):
    """
    Create a subdocument table with color-coded concordance cells.
    Columns: Facility | Province | District | Indicator | MRF | E-HR | Concordance (%)
    """
    sub = tpl.new_subdoc()
    tbl = sub.add_table(rows=1, cols=7)
    headers = ["Facility", "Province", "District", "Indicator", "MRF", "E-HR", "Concordance (%)"]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]; cell.text = h
        for run in cell.paragraphs[0].runs: run.bold = True

    # Sort for readability
    view = (table2_df
            .sort_values(["Province","District","Facility","Indicator"])
            .reset_index(drop=True))

    for _, r in view.iterrows():
        row = tbl.add_row().cells
        row[0].text = str(r["Facility"])
        row[1].text = str(r["Province"])
        row[2].text = str(r["District"])
        row[3].text = str(r["Indicator"])
        row[4].text = f"{int(round(r['MRF']))}"
        row[5].text = f"{int(round(r['E_HR']))}"
        conc = r["Concordance_%"]
        row[6].text = f"{conc:.1f}%" if not pd.isna(conc) else ""
        # Color-code the Concordance cell
        col = color_for_conc(conc)
        if col:
            shade_cell(row[6], col)
    return sub

def build_table3_subdoc(tpl, table3_df):
    """
    Create a subdocument summary table: Indicator | MRF Total | E-HR Total | Overall Concordance (%)
    """
    sub = tpl.new_subdoc()
    tbl = sub.add_table(rows=1, cols=4)
    headers = ["Indicator", "MRF Total", "E-HR Total", "Overall Concordance (%)"]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]; cell.text = h
        for run in cell.paragraphs[0].runs: run.bold = True

    # Show in the standard order
    order = ["HTS_TST", "HTS_POS", "TX_NEW", "TX_CURR"]
    view = (table3_df.set_index("Indicator").reindex(order).reset_index()
            .fillna({"MRF_Total":0,"E_HR_Total":0,"Overall_Concordance_%":np.nan}))
    for _, r in view.iterrows():
        row = tbl.add_row().cells
        row[0].text = str(r["Indicator"])
        row[1].text = f"{int(r['MRF_Total'])}"
        row[2].text = f"{int(r['E_HR_Total'])}"
        conc = r["Overall_Concordance_%"]
        row[3].text = f"{conc:.1f}%" if not pd.isna(conc) else ""
        col = color_for_conc(conc)
        if col:
            shade_cell(row[3], col)
    return sub

# -----------------------------
# Main
# -----------------------------
def main():
    ap = argparse.ArgumentParser(description="Generate CDC Concordance Tables & Figure and inject into template.")
    ap.add_argument("--excel", required=True, help="Path to Concordance_Analysis_Complete.xlsx")
    ap.add_argument("--template", default="cdc_template.docx", help="Path to cdc_template.docx")
    ap.add_argument("--outdir", default="output", help="Output directory")
    ap.add_argument("--period", default="COP24 Q4", help="Label used in titles/captions")
    ap.add_argument("--out", default=None, help="Output DOCX path")
    args = ap.parse_args()

    os.makedirs(args.outdir, exist_ok=True)
    stamp = datetime.now().strftime("%Y-%m-%d")

    # 1) Load Excel
    df_raw = load_concordance_excel(args.excel)

    # 2) Build Table 2 (facility-level)
    table2 = build_table2(df_raw)

    # 3) Build Table 3 (overall by indicator)
    table3 = build_table3(table2)

    # 4) Build Figure 4 (Harare vs Bulawayo)
    fig4_path = os.path.join(args.outdir, f"figure4_harare_bulawayo_{args.period.replace(' ', '')}_{stamp}.png")
    build_figure4(table2, fig4_path, args.period)

    # 5) Save CSVs for audit/reference
    t2_csv = os.path.join(args.outdir, f"table2_facility_concordance_{args.period.replace(' ', '')}_{stamp}.csv")
    table2.to_csv(t2_csv, index=False)
    t3_csv = os.path.join(args.outdir, f"table3_overall_concordance_{args.period.replace(' ', '')}_{stamp}.csv")
    table3.to_csv(t3_csv, index=False)

    # 6) Render into Word template
    tpl = DocxTemplate(args.template)
    sub_table2 = build_table2_subdoc(tpl, table2)
    sub_table3 = build_table3_subdoc(tpl, table3)

    context = {
        "table2_concordance_facility": sub_table2,
        "table3_overall_concordance": sub_table3,
        "figure4": InlineImage(tpl, fig4_path, width=Mm(140)),
    }
    tpl.render(context)

    out_docx = args.out or os.path.join(args.outdir, f"CDC_Concordance_{args.period.replace(' ', '')}_{stamp}.docx")
    tpl.save(out_docx)

    print("Done.")
    print("Table 2 CSV:", t2_csv)
    print("Table 3 CSV:", t3_csv)
    print("Figure 4  :", fig4_path)
    print("Rendered  :", out_docx)

if __name__ == "__main__":
    main()
